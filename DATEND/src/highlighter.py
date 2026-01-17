import docx
from docx.enum.text import WD_COLOR_INDEX
import os
import storage
from PySide6.QtWidgets import QFileDialog, QApplication
from docx.shared import RGBColor
import re
import pickle
import random
import time

def highlighter(filepath):
    canpredict = True
    with open('richpredictor.pickle', 'rb') as file:
        richpredict = pickle.load(file)
    print('imported model')
    
    literarydevices = 0
    ar = 0
    evidences = 0
    transitionwords = 0
    richideas = 0
    doc = docx.Document(filepath)
    print('imported')
    devices = []
    orange = RGBColor(255, 153, 0)
    purple = RGBColor(218, 199, 229)
    did = False

    devlist = open('litdevices.txt', 'r')
    
    for device in devlist:
        delist = device.split('\n')
        for device in delist:
            if device != '':
                print(device)
                if device.endswith(','):
                    device.strip(',')
                devices.append(device)
        
    author = ['author', 'reader', 'playwrite', 'audience']
    for name in storage.author_names:
        print(name)
        names = name.split()
        for nm in names:
            if nm.endswith(','):
                nm.strip(',')
            if "'" in nm:
                altered = nm.replace("'", "’")
                author.append(altered)
            
            author.append(nm)
    print(author)
    evidence = False
    tobefalse = False
    transitions = ['Furthermore,', 'Moreover,', 'Additionally,', 'Besides,', 'However,', 'Nonetheless,', 'Nevertheless,', 'Conversely,', 'Similarly,', 'Likewise,', 'Consequently,', 'Therefore,', 'Thus,', 'Hence,', 'Specifically,', 'Overall,', 'Ultimately,']
    #for rich idea purposes:
    hasauthor = 0
    hasfeatures = 0
    hasthat = 0
    hasconvey = 0
    hastheme = 0
    hasissue = 0
    long = 0
    hasproblem = 0
    hasquestion = 0
    #actual related data for rich ideas:
    questionwords = storage.gq_words
    themes = storage.themes
    issues = storage.global_issues
    synproblems = ['issue', 'problem', 'dispute', 'obstacle', 'conundrum', 'enigma', 'question', 'difficult','challeng']
    synconvey = ['highlight', 'communicat', 'express', 'convey', 'reveal', 'tell', 'say', 'show', 'emphasis', 'emphasiz', 'identif', 'accentuat', 'stress', 'punctuat', 'spotlight', 'pinpoint', 'illuminat', 'describ', 'show', 
                 'discuss', 'support', 'embod']

    istransition = storage.transition_words
    isTextual = storage.textual
    isRich = storage.rich
    isstats = storage.stats
    richness = False
    highlightrich = False

    new_doc = docx.Document()
    for para in doc.paragraphs:
            new_para = new_doc.add_paragraph()
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            for run in para.runs:
                words = run.text.split()
                for word in words:
                    did = False
                    #literary devices: Yellow highlights
                    for device in devices:
                        if device.lower() == word.lower() and did == False and evidence == False:
                            did = True
                            literarydevices += 1
                            new_run = new_para.add_run(word)
                            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            default_settings(new_run, run)
                            new_run = new_para.add_run(' ')
                            default_settings(new_run, run)
                            break

                    #transition words = purple higlights
                
                    if istransition:
                        for transition in transitions:
                            if transition.lower() == word.lower() and did == False and evidence == False:
                                did = True
                                transitionwords += 1
                                new_run = new_para.add_run(word)
                                #change font color
                                new_run.font.color.rgb = RGBColor(128, 0, 128)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                new_run.font.size = run.font.size
                                new_run.font.name = run.font.name
                                new_run = new_para.add_run(' ')
                                default_settings(new_run, run)
                                break
    

                    #author/reader = orange highlights
                    for authorreader in author:
                        if authorreader.lower() == word.lower() and did == False and evidence == False:
                            did = True
                            ar += 1
                            new_run = new_para.add_run(word)
                            new_run.font.highlight_color = WD_COLOR_INDEX.DARK_YELLOW
                            if word in run.text and run.font:
                                new_run.font.highlight_color.rgb = orange
                            default_settings(new_run, run)
                            new_run = new_para.add_run(' ')
                            default_settings(new_run, run)
                            break

                    #textual evidence =  blue highlights
                    if isTextual:
                        if '"' in word or '“' in word or "‘" in word:
                            evidences += 1
                            evidence = True
                        if ('"' in word or '”' in word or "’" in word) and evidence == True:
                            tobefalse = True

                    
                    if evidence == True:
                        new_run = new_para.add_run(word)
                        new_run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        default_settings(new_run, run)
                        new_run = new_para.add_run(' ')
                        if tobefalse == False:
                            new_run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        default_settings(new_run, run)

                        if tobefalse == True:
                            evidence = False
                            tobefalse = False
                    else:
                        if did == False:
                            new_run = new_para.add_run(word + ' ')
                            new_run.font.highlight_color = run.font.highlight_color
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = run.font.size
                            new_run.font.name = run.font.name
                            new_run.font.color.rgb = run.font.color.rgb
            
            new_para.add_run().add_break()
    if isstats:
        print('add stats')
        new_doc.add_paragraph('Statistics Report:\nLiterary devices used: ' + str(literarydevices) + 
                              '\nAuthor-Reader Relationship: ' + str(ar) + 
                              '\nTextual evidences used: ' + str(evidences) + 
                              '\nTransition words used: ' + str(transitionwords))
    
    #main stuff rich ideas
    if isRich and canpredict:  
        temp_doc = docx.Document()
        for paras in new_doc.paragraphs: 
            runs_to_combine = []
            new_paragraph = temp_doc.add_paragraph()
            new_paragraph.style = paras.style
            new_paragraph.alignment = paras.alignment
            for run in paras.runs:
                word = str(run.text)
                word.strip('')
                if not word.endswith('.') and not word.endswith('?') and not word.endswith('!') and word != '\n':
                    runs_to_combine.append(word)
                else:
                    runs_to_combine.append(word)
                    sentence_add = ' '.join(runs_to_combine)
                    print(sentence_add)
                    sentence_run = new_paragraph.add_run(sentence_add + ' ')
                    sentence_run.bold = run.bold
                    sentence_run.italic = run.italic
                    sentence_run.underline = run.underline
                    sentence_run.font.size = run.font.size
                    sentence_run.font.name = run.font.name
                    sentence_run.font.color.rgb = run.font.color.rgb 
                    runs_to_combine = []

        brand_new_doc = docx.Document()
        for paragraph in temp_doc.paragraphs:
            brand_new_para = brand_new_doc.add_paragraph()
            brand_new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            brand_new_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
            brand_new_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
            for run in paragraph.runs:
                print(run.text)
                sentences = run.text.split('. ')
                print(sentences)
                for sentence in sentences:
                    number_words = sentence.split()
                    if len(number_words) >= 18:
                        long = 1
                        #print(len(sentence))
                    for word in questionwords:
                        if search(word, sentence):
                            hasquestion = 1
                            #print('has question')
                    for word in themes:
                        if search(word, sentence):
                            hastheme = 1
                            #print('has theme')
                    if search('that', sentence):
                        hasthat = 1
                        #print('that found')
                    for word in issues:
                        if search(word, sentence):
                            hasissue = 1
                            #print('has issue')
                    for word in synproblems:
                        if search(word, sentence):
                            hasproblem = 1
                            #print('has problem')
                    for word in synconvey:
                        if search(word, sentence):
                            hasconvey = 1
                            #print('has convey')
                    for word in author:
                        if search(word, sentence):
                            #print('has author')
                            hasauthor = 1
                    for word in devices:
                        if search(word, sentence):
                            hasfeatures = 1
                            #print('has features')
                    topredict = [[hasauthor, hasfeatures, hasthat, hasconvey, hastheme, hasissue, long, hasproblem, hasquestion]]

                    prob_yes = richpredict.predict_proba(topredict)[0][1]
                    
                    if prob_yes >= 0.95:
                        highlightrich = True
                    elif prob_yes >= 0.8 and prob_yes < 0.95:
                        lottery = random.random()
                        if lottery <= prob_yes:
                            highlightrich = True

                    '''
                    if richpredict.predict(topredict):
                        highlightrich = True
                    '''
                    if highlightrich:
                        print('rich found')
                        richideas += 1
                        brand_new_run = brand_new_para.add_run()
                        brand_new_run.text = sentence
                        print(sentence)
                        brand_new_run.font.highlight_color = WD_COLOR_INDEX.PINK
                        default_settings(brand_new_run, run)
                        brand_new_run = brand_new_para.add_run('. ')
                        default_settings(brand_new_run, run)
                    else:
                        brand_new_run = brand_new_para.add_run()
                        brand_new_run.text = sentence
                        default_settings(brand_new_run, run)
                        brand_new_run.font.highlight_color = run.font.highlight_color
                        brand_new_run = brand_new_para.add_run('. ')
                        default_settings(brand_new_run, run)

                    #reset
                    highlightrich = False   
                    hasauthor = 0
                    hasfeatures = 0
                    hasthat = 0
                    hasconvey = 0
                    hastheme = 0
                    hasissue = 0
                    long = 0
                    hasproblem = 0
                    hasquestion = 0


    if isstats:
        print('add stats')
        brand_new_doc.add_paragraph('Statistics Report:\nLiterary devices used: ' + str(literarydevices) + 
                              '\nAuthor-Reader Relationship: ' + str(ar) + 
                              '\nTextual evidences used: ' + str(evidences) + 
                              '\nTransition words used: ' + str(transitionwords) + 
                              '\nRich Ideas present: ' + str(richideas))
    

    folder_path = QFileDialog.getExistingDirectory(None, "Select Folder to download")
    if folder_path:
        file_filter = "Word Documents (*.docx)"
        filename, _ = QFileDialog.getSaveFileName(None, "Save As", filter=file_filter)
        if filename:
            full = filename
            print(full)
            new_doc.save(full) 
            print('saved')
    
    if isRich and canpredict:
        folder_path = QFileDialog.getExistingDirectory(None, "Select Folder to download")
        if folder_path:
            file_filter = "Word Documents (*.docx)"
            filename, _ = QFileDialog.getSaveFileName(None, "Save As", filter=file_filter)
            if filename:
                full = filename
                print(full)
                brand_new_doc.save(full) 
                print('saved')


def default_settings(torun, run):
    torun.bold = run.bold
    torun.italic = run.italic
    torun.underline = run.underline
    torun.font.size = run.font.size
    torun.font.name = run.font.name
    torun.font.color.rgb = run.font.color.rgb

def search(voc, sen):
    pattern = r'\b{}[a-zA-Z]*\b'.format(voc)
    return re.search(pattern, sen)

'''
app = QApplication([])
file_dialog = QFileDialog()
file_dialog.setFileMode(QFileDialog.ExistingFile)

if file_dialog.exec() == QFileDialog.Accepted:
    selected_file = file_dialog.selectedFiles()[0]
highlighter(selected_file)
'''

    
