import docx
import re
from docx.enum.text import WD_COLOR_INDEX


word_list = []
csv = open('list.csv', 'r')
for line in csv:
    line = line.strip()
    word_list.append(line)

doc = docx.Document("sample.docx")
highlighted_doc = docx.Document()
highlight_color = WD_COLOR_INDEX.YELLOW

for paragraph in doc.paragraphs:
    new_paragraph = highlighted_doc.add_paragraph()
    words = paragraph.text.split()

    for word in words:
        if word.strip(",.") in word_list:
            for run in paragraph.runs:
                if word in run.text:
                    new_run = new_paragraph.add_run(run.text.replace(word, f"{word}", 1))
                    new_run.font.highlight_color = highlight_color
        else:
            new_paragraph.add_run(word + " ")


    if not new_paragraph.runs:
        new_paragraph.add_run(paragraph.text)


highlighted_doc.save("highlighted.docx")
